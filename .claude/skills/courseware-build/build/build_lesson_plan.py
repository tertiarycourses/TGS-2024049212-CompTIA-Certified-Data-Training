#!/usr/bin/env python3
"""Generate the WSQ CompTIA Certified Data+ Lesson Plan (LP) DOCX in the Tertiary house format.

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 5-day schedule tables (9:00am-6:00pm, 8 training hours/day, 1h
lunch, tea within, final assessment Day 5 4:00pm-6:00pm). Topics/labs come from
course_data + the domain data files so the LP stays aligned with the deck,
guide and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"

import json as _json
_SM_PATH=os.path.join(HERE,"slide_map.json")
SLIDE_MAP=_json.load(open(_SM_PATH)) if os.path.exists(_SM_PATH) else {"domains":{},"labs":{}}

def slides_for(labs=None, domain=None):
    """Slide range cited in the Lesson Plan, read from the deck's exported slide map."""
    if labs:
        rs=[SLIDE_MAP["labs"].get(str(n)) for n in labs]
        rs=[r for r in rs if r]
        if rs: return f"{min(r[0] for r in rs)}–{max(r[1] for r in rs)}"
    if domain:
        r=SLIDE_MAP["domains"].get(str(domain))
        if r: return f"{r[0]}–{r[1]}"
    return "—"

def lab_titles(nums):
    # a['title'] already begins with "Lab N — ", so don't prefix it again
    return "; ".join(a['title'] for a in ACT if a['num'] in nums)

# ------------------------------------------------ schedule (single source of truth for timing)
# (start, end, minutes, kind, activity_text)  kind: admin/topic/lab/break/lunch/assess/recap
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:00","9:30",30,"admin","Welcome, trainer and learner introductions, ground rules, learning outcomes, course outline and mandatory digital attendance (AM)","1–19"),
    ("9:30","10:30",60,"topic","Domain 1 — Data Concepts: data schemas, relational vs non-relational databases, the four NoSQL families, normalisation 1NF-5NF (concepts + demo)",slides_for(domain=1)),
    ("10:30","10:45",15,"break","Tea break","—"),
    ("10:45","13:00",135,"lab","Domain 1 — keys, relationships, referential integrity, OLTP vs OLAP. Hands-on: "+lab_titles([1]),slides_for(labs=[1])),
    ("13:00","14:00",60,"lunch","Lunch break","—"),
    ("14:00","15:30",90,"lab","Domain 1 — warehouses, marts, lakes and lakehouses, star vs snowflake schemas, slowly changing dimensions, data types and structures. Hands-on: "+lab_titles([2]),slides_for(labs=[2])),
    ("15:30","15:45",15,"break","Tea break","—"),
    ("15:45","17:45",120,"lab","Domain 1 — file formats, data languages, infrastructure (cloud/on-premise/containers), AI concepts and data sources. Hands-on: "+lab_titles([3]),slides_for(labs=[3])),
    ("17:45","18:00",15,"recap","Day 1 recap, Q&A and PM digital attendance","1–19"),
 ]),
 2: (C.DAY_THEMES[2], [
    ("9:00","9:15",15,"recap","Day 1 recap and mandatory digital attendance (AM)","1–19"),
    ("9:15","10:30",75,"topic","Domain 2 — Data Acquisition: ETL vs ELT, full and delta loads, APIs (pull/push), web scraping, public repositories, surveys and sampling (concepts + demo)",slides_for(domain=2)),
    ("10:30","10:45",15,"break","Tea break","—"),
    ("10:45","13:00",135,"lab","Domain 2 — data profiling procedure, the defect classes: missing values, duplicates, redundancy, invalid data and outliers. Hands-on: "+lab_titles([4]),slides_for(labs=[4])),
    ("13:00","14:00",60,"lunch","Lunch break","—"),
    ("14:00","15:30",90,"lab","Domain 2 — data transformation: recoding, derived variables, imputation, aggregation, transposing, appending, parsing strings. Hands-on: "+lab_titles([5]),slides_for(labs=[5])),
    ("15:30","15:45",15,"break","Tea break","—"),
    ("15:45","17:45",120,"lab","Domain 2 — text/date/logical functions, SQL joins, filtering, indexing, subqueries and query execution plans. Hands-on: "+lab_titles([6,7]),slides_for(labs=[6, 7])),
    ("17:45","18:00",15,"recap","Day 2 recap, Q&A and PM digital attendance","1–19"),
 ]),
 3: (C.DAY_THEMES[3], [
    ("9:00","9:15",15,"recap","Day 2 recap and mandatory digital attendance (AM)","1–19"),
    ("9:15","10:30",75,"topic","Domain 3 — analysis types (exploratory, performance, trend, gap, link) and descriptive statistics: mean, median, mode, range, variance, standard deviation (concepts + demo)",slides_for(domain=3)),
    ("10:30","10:45",15,"break","Tea break","—"),
    ("10:45","13:00",135,"lab","Domain 3 — z-scores, normal distribution, the empirical rule, parametric vs nonparametric data, frequency and histograms. Hands-on: "+lab_titles([8]),slides_for(labs=[8])),
    ("13:00","14:00",60,"lunch","Lunch break","—"),
    ("14:00","15:30",90,"lab","Domain 3 — inferential statistics: hypothesis testing, t-tests, p-values, Type I and Type II errors, chi-square, confidence intervals. Hands-on: "+lab_titles([9]),slides_for(labs=[9])),
    ("15:30","15:45",15,"break","Tea break","—"),
    ("15:45","17:45",120,"lab","Domain 3 — correlation, regression, R-squared, causation, troubleshooting analysis issues and communicating results to different audiences. Hands-on: "+lab_titles([10]),slides_for(labs=[10])),
    ("17:45","18:00",15,"recap","Day 3 recap, Q&A and PM digital attendance","1–19"),
 ]),
 4: (C.DAY_THEMES[4], [
    ("9:00","9:15",15,"recap","Day 3 recap and mandatory digital attendance (AM)","1–19"),
    ("9:15","10:30",75,"topic","Domain 4 — chart selection by question type: composition, comparison, distribution, relationship and trend; line, bar, pie, histogram, scatter (concepts + demo)",slides_for(domain=4)),
    ("10:30","10:45",15,"break","Tea break","—"),
    ("10:45","13:00",135,"lab","Domain 4 — specialist charts (tree map, heat map, bubble, waterfall, combination), geographic maps: dot, filled and layered. Hands-on: "+lab_titles([11]),slides_for(labs=[11])),
    ("13:00","14:00",60,"lunch","Lunch break","—"),
    ("14:00","15:30",90,"lab","Domain 4 — expressing business requirements, report and dashboard components, documentation elements, report types. Hands-on: "+lab_titles([12]),slides_for(labs=[12])),
    ("15:30","15:45",15,"break","Tea break","—"),
    ("15:45","17:45",120,"lab","Domain 4 — dashboard design principles, filtering and sorting, drillthrough and tooltips, validating reporting accuracy, validation vs verification. Dashboard build review and critique",slides_for(domain=4)),
    ("17:45","18:00",15,"recap","Day 4 recap, Q&A and PM digital attendance","1–19"),
 ]),
 5: (C.DAY_THEMES[5], [
    ("9:00","9:15",15,"recap","Day 4 recap and mandatory digital attendance (AM)","1–19"),
    ("9:15","10:30",75,"topic","Domain 5 — data governance, the data lifecycle, data roles (owner, steward, custodian, privacy officer), classification and compliance (concepts + demo)",slides_for(domain=5)),
    ("10:30","10:45",15,"break","Tea break","—"),
    ("10:45","12:15",90,"lab","Domain 5 — privacy and protection: access control, encryption at rest/in transit/in use, masking, de-identification and pseudonymisation. Hands-on: "+lab_titles([13]),slides_for(labs=[13])),
    ("12:15","13:00",45,"lab","Domain 5 — retention, preservation, destruction and sanitisation; regulations (PDPA, GDPR, HIPAA, SOX) and data sovereignty. Hands-on: "+lab_titles([14]),slides_for(labs=[14])),
    ("13:00","14:00",60,"lunch","Lunch break","—"),
    ("14:00","15:15",75,"lab","Domain 5 — documentation, versioning and data lineage; the six quality dimensions; profiling, monitoring and automated testing; master data management. Hands-on: "+lab_titles([15]),slides_for(labs=[15])),
    ("15:15","15:25",10,"break","Tea break","—"),
    ("15:25","15:45",20,"recap","Course revision across all five exam domains, CompTIA Data+ certification exam guidance and Q&A",slides_for(domain=5)),
    ("15:45","15:55",10,"admin","Course feedback and mandatory TRAQOM survey on the LMS","1–19"),
    ("15:55","16:00",5,"assess","Briefing for Assessment and Assessment digital attendance","196–200"),
    ("16:00","17:00",60,"assess","Written Assessment (WA) — Short-Answer Questions (SAQ), 1 hour, individual, open book","196–200"),
    ("17:00","18:00",60,"assess","Practical Performance (PP) — scenario-based hands-on data tasks, 1 hour, individual, open book. PM digital attendance","196–200"),
 ]),
}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("4.0","1 September 2024","Previous release — CompTIA Data+ course slides and lesson plan (legacy master deck).",C.TRAINER),
 ("5.0",C.VERSION_DATE,"Full rebuild against the CompTIA Data+ (DA0-001) exam domains and the approved Course Proposal (LU1-LU5 / LO1-LO5, TSC ATP-PIN-3001-1.1). 5-day / 40-hour schedule with 15 hands-on labs; assessment block moved to Day 5, 4:00-6:00 pm.",C.TRAINER),
 ("5.1",C.VERSION_DATE,"Each hands-on lab now ships its own dataset (CSV, plus XLSX where a spreadsheet is the natural tool) under labs/<lab>/data/, replacing the typed-in sample data. Datasets enlarged to realistic sample sizes so the statistics behave correctly; every expected result re-verified by executing the labs against the shipped data. Lab 1 DDL, Lab 12 fault injection and Lab 8 z-threshold corrected.",C.TRAINER),
])
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration",f"{C.DAYS} days · 8 training hours per day ({C.DAYS*8} hours) — {C.CLASSROOM_HOURS} h classroom facilitation + {C.PRACTICAL_HOURS} h practical + {C.ASSESSMENT_HOURS} h assessment = {C.TOTAL_HOURS} hours"),
      ("Daily Timing","9:00 am – 6:00 pm (1-hour lunch; tea breaks within training time)"),
      ("Certification Alignment",f"{C.CERT_EXAM} — all five exam domains at their examined weightings"),
      ("Skills Framework",f"TSC: {C.TSC_TITLE} ({C.TSC_CODE})"),
      ("Mode","Instructor-led classroom facilitation with hands-on data labs in the browser (Killercoda Ubuntu, SQLite, Python/pandas) and browser-based analysis tools"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "Final assessment is conducted on Day 5 from 4:00 pm to 6:00 pm (WA 1 hour followed by PP 1 hour).",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":"F3F5F8","recap":"F3F5F8","lab":None}

H("Course Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=4); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity","Slides"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0
    for start,end,mins,kind,text,slides in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[3],slides if kind not in ("break","lunch") else "—",size=9.5,fill=fill)
        if kind!="lunch": training+=mins
    # widths
    for row in tbl.rows:
        row.cells[0].width=Inches(1.0); row.cells[1].width=Inches(0.75)
        row.cells[2].width=Inches(4.15); row.cells[3].width=Inches(0.85)
    p=doc.add_paragraph(); r=p.add_run(f"Total training time: {training} minutes ({training//60} hours)."); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    assert training==480, f"Day {day} training minutes = {training}, expected 480"

p=doc.add_paragraph(); r=p.add_run(
 f"Slide numbers refer to the Trainer Slides deck {C.SHORT_TITLE}-{C.VERSION}. "
 "Any change to the deck requires this column to be re-checked.")
r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY

H("Lab Reference (aligned to the CompTIA Data+ exam domains)",1)
tt=doc.add_table(rows=0,cols=4); tt.style="Table Grid"
hdr=tt.add_row().cells
for i,htext in enumerate(["Exam domain / Learning Unit","Exam weighting","Labs","Slides"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for tp in C.TOPICS:
    acts=[a for a in ACT if a["topic"]==tp["num"]]
    cells=tt.add_row().cells
    set_cell(cells[0],f"Domain {tp['code']} ({tp['lu']} / {tp['lo']}): {tp['title']}",bold=True,size=9.5,fill=TOPIC_FILL)
    set_cell(cells[1],tp["weighting"],size=9.5,fill=TOPIC_FILL)
    set_cell(cells[2],", ".join(f"Lab {a['num']}" for a in acts),size=9.5)
    set_cell(cells[3],slides_for(domain=tp["num"]),size=9.5)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)
