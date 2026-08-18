#!/usr/bin/env python3
"""Generate the WSQ CompTIA Certified Data+ Learner Guide as BOTH a Markdown mirror (LG-*.md at repo
root) and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per lab (Objective · Goal · What you'll build · Step-by-step
with commands · Test it), plus setup, exam-prep and glossary. All content is
driven by course_data + the domain data files, keeping the LG 100% aligned with
the slide deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))
def image(path,caption=""): B.append(("image",path,caption))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  f"It provides detailed step-by-step instructions for all {len(ACT)} hands-on labs, organised by the five "
  f"official {C.CERT_EXAM} exam domains. Every lab maps to a published exam objective and to the course "
  "Learning Outcomes (LO1-LO5) approved under the Skills Framework TSC "
  f"{C.TSC_TITLE} ({C.TSC_CODE}).")
p("Use this guide alongside the course slides and the lab files in the labs/ folder of the course repository. "
  "The slides carry the concepts and the visual explanations; this guide carries the full click-by-click and "
  "command-by-command procedure for every lab, together with the expected result and a troubleshooting table. "
  "Work through the labs in order — each one builds on a technique established in the labs before it.")
p("All labs run free in your browser. The terminal-based labs use the Killercoda Ubuntu playground "
  "(https://killercoda.com/playgrounds/scenario/ubuntu) with Python 3, pandas and SQLite; the remaining labs "
  "use purpose-built browser tools that parse everything locally on your own machine, so no data you load "
  "ever leaves your computer.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Before You Start — Environment Setup")
h3("What you need")
bullets([
 "A modern web browser (Chrome, Edge, Firefox or Safari) and an internet connection. Nothing needs to be installed on your own machine.",
 "The Killercoda Ubuntu playground — https://killercoda.com/playgrounds/scenario/ubuntu — a free browser-based Ubuntu terminal used by the terminal labs. A free account keeps your session alive longer.",
 "The four browser-based course tools listed below, each of which parses its input locally in your browser.",
 "The course lab repository — clone it, or download it as a ZIP from GitHub (link on your LMS course page).",
])
h3("The browser-based tools used in the labs")
bullets([f"{name} — {desc}  ({url})" for name, url, desc in C.LAB_TOOLS])
h3("Verify the Killercoda environment")
p("Open the Killercoda Ubuntu playground and confirm the three tools the terminal labs depend on are present. "
  "Killercoda gives you a throwaway Ubuntu machine in the browser: anything you create is discarded when the "
  "session ends, so download any file you want to keep before you close the tab.")
code("$ python3 --version          # Python 3 is pre-installed\n"
     "$ sqlite3 --version          # SQLite 3 is pre-installed\n"
     "$ pip3 install pandas matplotlib scipy --quiet   # the analysis libraries used from Lab 4 onwards\n"
     "$ python3 -c \"import pandas, matplotlib, scipy; print('all libraries ready')\"")
note("If pip3 refuses to install because the environment is externally managed, add the "
     "--break-system-packages flag: pip3 install pandas --quiet --break-system-packages. "
     "This is safe on a throwaway Killercoda instance.")
h3("Conventions used in every lab")
bullets([
 "Commands are shown with a leading $ prompt. Type or paste everything after the $, not the $ itself.",
 "Each lab creates its own working directory under ~/dataplus/ so the labs never overwrite one another.",
 "Multi-line files are created with a heredoc (cat > file <<'EOF' ... EOF). Paste the whole block at once; if your browser mangles it, use nano instead and paste the contents.",
 "Every lab ends with a 'Test it' check stating the exact expected result — do not move on until yours matches.",
 "Each lab also carries a troubleshooting table covering the three failures learners hit most often.",
 "Download any CSV, PNG or report you want to keep before your Killercoda session expires.",
])

# ---------------- per-topic, per-lab ----------------
TOPICS_BY_NUM={t["num"]:t for t in C.TOPICS}
for t in C.TOPICS:
    h1(f"Domain {t['code']} — {t['title']}  ({t['weighting']} of the exam)")
    p(t["subtitle"])
    h3("Key concepts")
    bullets(t["concepts"])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(a["title"])   # already formatted "Lab N — Title"
        p(f"Exam objective: {a['objective']}")
        p(f"Where you run it: {a.get('env','Killercoda Ubuntu playground')}")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Tools: {a['services']}.)")
        if a.get("figure"):
            image(os.path.join(REPO,"courseware","assets",a["figure"]),
                  f"Expected output for {a['title'].split('—')[0].strip()} — your run should match this.")
        h3("Step-by-step")
        st=[]
        for i,(instr,cmd) in enumerate(a["steps"],1):
            st.append((instr,cmd))
        steps(st)
        h3("Test it — the expected result")
        p(a["test"])
        if a.get("troubleshoot"):
            h3("If it doesn't work")
            bullets([f"{sym} — {fix}" for sym, fix in a["troubleshoot"]])
        note(f"The same procedure, with the copy-paste commands, is in labs/lab-{a['num']:02d}/README.md "
             f"in the course repository.")
        rule()

h1("Exam Focus — Cross-Cutting Topics")
p(f"These topics are examined across the {C.CERT_EXAM} blueprint but do not each carry a dedicated hands-on "
  "lab. Study this section alongside the labs so you can answer the knowledge questions confidently.")
h3("Data acquisition concepts you must be able to compare")
bullets([
 "ETL vs ELT — ETL transforms before loading (schema-on-write, warehouse); ELT loads raw then transforms (schema-on-read, lake/lakehouse).",
 "Full load vs delta (incremental) load — a full load reloads everything and self-corrects; a delta load is fast but depends on a reliable watermark.",
 "API pull vs push — a pull model polls on your schedule; a push model (webhook) notifies you on change and is closer to real time.",
 "Synchronous vs asynchronous web services — synchronous waits for the response; asynchronous lets you continue working while it completes.",
 "Sampling methods — simple random, systematic (every nth record) and stratified (proportional within defined groups).",
 "Survey design — bias is the enemy; use a Likert scale for shades of opinion, and make sure the response options cover the full range.",
])
h3("Statistical terms you must be able to define")
bullets([
 "Population vs sample — the whole group vs the subset you actually measured. Inferential statistics generalise from one to the other.",
 "Parametric vs nonparametric data — parametric fits a known (usually normal) distribution; nonparametric does not, so use distribution-independent methods.",
 "The empirical rule — in a normal distribution roughly 68%, 95% and 99.74% of values lie within one, two and three standard deviations of the mean.",
 "Statistical significance vs practical significance — p < 0.05 says the effect is unlikely to be chance; it does NOT say the effect is large enough to matter commercially.",
 "Type I error (false positive) — rejecting a true null hypothesis. Type II error (false negative) — failing to reject a false null hypothesis.",
 "Confidence interval — the range within which the true population value is likely to lie, at a stated confidence level (commonly 95%).",
 "R and R-squared — Pearson's r measures the strength and direction of a linear relationship; R-squared is the proportion of variance explained.",
])
h3("Data quality dimensions (examined in Domain 5)")
bullets([
 "Completeness — no required value is missing.",
 "Accuracy — the values are correct and match the real-world fact.",
 "Consistency — the same thing is recorded the same way in every system.",
 "Uniqueness — no unintended duplicate records; the primary key identifies exactly one row.",
 "Validity — values conform to the defined type, format, range and domain.",
 "Timeliness — the data is current enough for the decision it supports.",
])
h3("Governance vocabulary")
bullets([
 "Data owner (senior accountability) vs data steward (labelling and quality) vs data custodian (the systems) vs privacy officer (privacy oversight).",
 "Classification (public, internal, sensitive, confidential, restricted) vs data type tag (PII, PHI, PIFI, intellectual property).",
 "Masking (hides part of a value) vs de-identification (removes identifiers) vs pseudonymisation (substitutes a surrogate key, reversible only with the mapping).",
 "Data at rest, in transit and in use — the three states, each needing its own protection.",
 "Retention (how long you keep it) vs preservation (a hold outside the retention policy) vs removal, destruction and sanitisation (which verifies the wipe).",
 "Data lineage (the traceable path from source to report) and the data dictionary (the shared definition of every field).",
 "Master data management — maintaining the 'golden record' for core entities so every system agrees.",
])
rule()

h1("Exam Preparation")
bullets([
 "First pass: complete every lab in order, reading the concept slides for that domain before you start.",
 "Second pass: redo each lab from a blank terminal until the workflow is automatic without the guide.",
 "Review the 'Test it' expected result for every lab — if you cannot predict it, re-read that domain.",
 "Know the exam weightings and revise proportionally: Data Analysis 24%, Data Acquisition and Preparation 22%, "
 "Data Concepts and Environments 20%, Visualization and Reporting 20%, Data Governance 14%.",
 "Practise the comparisons the exam loves: relational vs non-relational, OLTP vs OLAP, ETL vs ELT, "
 "star vs snowflake, mean vs median, validation vs verification, masking vs de-identification.",
 "Be able to choose the correct chart from the question asked — this is heavily tested in Domain 4.",
 "Sharpen exam readiness with the Tertiary Infotech CompTIA Data+ practice exams: {}".format(C.PRACTICE_EXAM_URL),
 "Take the free CompTIA practice assessment for DA0-001 and sit the exam via a Pearson VUE test centre "
 "or online proctoring.",
])

h1("Glossary")
gl=[
 ("Schema","A description of how data is organised and how tables relate to one another."),
 ("Primary key","A unique, non-null identifier for a record. Every table needs one."),
 ("Foreign key","A primary key from another table, referenced to create a relationship."),
 ("Referential integrity","The guarantee that every foreign key value exists in the parent table — no orphaned records."),
 ("Normalisation","Organising data (1NF-5NF) to remove redundancy so an update happens in exactly one place."),
 ("Denormalisation","Deliberately re-introducing redundancy to avoid expensive joins and speed up analytical queries."),
 ("OLTP / OLAP","Online Transactional Processing (many small real-time transactions) / Online Analytical Processing (long complex queries)."),
 ("Data warehouse","A combined, structured store of data from many source systems — the single source of truth."),
 ("Data mart","A subset of the warehouse serving one department or group."),
 ("Data lake / lakehouse","A cheap repository holding structured and unstructured data (schema-on-read); a lakehouse adds a schema layer so it can be queried like a warehouse."),
 ("Fact / dimension table","The fact table holds the measures and keys; dimension tables hold the descriptive attributes."),
 ("Star / snowflake schema","One ring of denormalised dimensions around a fact table; or dimensions that branch into further normalised dimensions."),
 ("Slowly changing dimension","How dimension changes are handled: Type 1 overwrites, Type 2 keeps full history, Type 3 keeps current and previous."),
 ("ETL / ELT","Extract-Transform-Load (transform before landing) / Extract-Load-Transform (land raw, transform later)."),
 ("Full / delta load","Reloading all data every run, versus loading only new or changed records."),
 ("API","Application Programming Interface — a defined request/response contract between systems. Pull polls; push notifies."),
 ("Data profiling","The first disciplined pass over a dataset: source, fields, types, keys and defect counts."),
 ("Null","A missing value, shown as blank, NULL or N/A. Not the same as zero or an empty string."),
 ("Outlier","A value far outside the normal distance from the rest of the data; commonly flagged at |z| > 3."),
 ("Derived variable","A new field computed from existing fields — a named data-manipulation technique."),
 ("Imputation","Substituting a missing value with an estimate such as the mean, median or a modelled value."),
 ("Join","Combining tables on a shared key: inner (matches only), left/right outer (keeps one side), full outer (keeps both)."),
 ("Index","A structure that speeds up lookups on a column, at the cost of space and slower writes."),
 ("Mean / median / mode","The arithmetic average / the middle value when sorted / the most frequent value."),
 ("Variance / standard deviation","The average squared deviation from the mean / its square root, in the original units."),
 ("Z-score","How many standard deviations a value lies from the mean: z = (x − x̄) / s."),
 ("Normal distribution","The symmetric bell curve; the empirical rule puts ~99.74% of values within three standard deviations."),
 ("Null / alternative hypothesis","H0 assumes no relationship between the variables; H1 assumes a relationship exists."),
 ("p-value","The probability that an observed difference arose by chance. Below 0.05 is conventionally significant."),
 ("Type I / Type II error","Rejecting a true null (false alarm) / failing to reject a false null (missed effect)."),
 ("Correlation (r) / R-squared","The strength and direction of a linear relationship / the proportion of variance explained."),
 ("Regression","A method estimating the relationship between a dependent and one or more independent variables."),
 ("KPI","Key Performance Indicator — a measure tied to a business objective."),
 ("Dashboard","A tool that tracks, analyses and displays data to support a small set of recurring decisions."),
 ("Validation / verification","Checking the format and structure of data / checking that the data is actually accurate."),
 ("Data governance","The organisational capability ensuring high-quality, controlled data across its whole lifecycle."),
 ("Data owner / steward / custodian","Senior accountability / labelling and quality / the systems the data lives on."),
 ("Data classification","Categorising data by sensitivity — public, internal, sensitive, confidential, restricted."),
 ("PII / PHI / PIFI","Personally Identifiable Information / Protected Health Information / Personally Identifiable Financial Information."),
 ("Masking / de-identification / pseudonymisation","Hiding part of a value / removing identifying fields / replacing an identifier with a surrogate key."),
 ("Data at rest / in transit / in use","Stored on disk / moving between systems / being processed in memory."),
 ("Retention / destruction / sanitisation","How long data is kept / deleting it and its medium / verifying the wipe was effective."),
 ("Data lineage","The traceable path from a data source through every transformation to the final report."),
 ("Data dictionary","The shared reference defining every field: meaning, type, valid values, owner and source."),
 ("Master data management","Maintaining the authoritative 'golden record' for core entities across all systems."),
 ("PDPA / GDPR","Singapore's Personal Data Protection Act / the EU General Data Protection Regulation."),
]
B.append(("dl",gl))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="image":
            rel=os.path.relpath(rest[0],REPO)
            out+=[f"![{rest[1]}]({rel})","",f"*{rest[1]}*",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("4.0","1 September 2024","Previous release — CompTIA Data+ Learner Guide (legacy master courseware).",C.TRAINER),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,f"Full rebuild against the {C.CERT_EXAM} exam domains and the approved Course Proposal (LU1-LU5 / LO1-LO5, TSC ATP-PIN-3001-1.1). {len(ACT)} browser-based hands-on labs with full step-by-step procedures, expected results and troubleshooting tables; new cross-cutting Exam Focus section and expanded data-analytics glossary.",C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,cmd) in enumerate(rest[0],1):
            para=doc.add_paragraph(style="List Number"); para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="image":
        if os.path.exists(rest[0]):
            from docx.shared import Inches as _In
            ip=doc.add_paragraph(); ip.alignment=WD_ALIGN_PARAGRAPH.CENTER
            ip.add_run().add_picture(rest[0], width=_In(6.0))
            cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cr=cp.add_run(rest[1]); cr.italic=True; cr.font.size=Pt(9); cr.font.color.rgb=GREY
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
