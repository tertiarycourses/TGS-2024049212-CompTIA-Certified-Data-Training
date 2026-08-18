#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Application Integration with Docker and Kubernetes' (TGS-2021010366):
  - Written Assessment (SAQ)  — 5 open-ended KNOWLEDGE questions (K1–K5), aligned to the slides
  - Practical Performance (PP) — 4 PRACTICAL tasks (LO1–LO4), aligned to the in-class activities
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Page 1 is the cover;
page 2 carries Trainee Information + Instructions + Grading; the questions/tasks begin on page 3.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "WSQ - CompTIA Certified Data+ Training"
COURSE_CODE = "TGS-2024049212"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = _logo("comptia-dataplus-badge.png")   # None if absent → Tertiary-only cover (as LP/LG)

Q_VER, A_VER = "v5", "v5"   # single standardised version across all four files
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# (criterion, context, question, [model-answer points]) — each traces to the course slides.
WRITTEN = [
 ("K1",
  "Before any analysis can be trusted, an analyst must know what defects the dataset carries. Data profiling is "
  "the disciplined first pass over a new dataset, and it is followed by the cleansing techniques that fix what "
  "the profiling found.",
  "Describe the process of data profiling and explain the tools and techniques used to identify and cleanse the "
  "main data-quality defects found in a raw dataset.",
  ["Data profiling is the process of examining a dataset to discover what information and trends it actually "
   "contains, BEFORE it is analysed or reported on.",
   "The profiling procedure: identify and document the source of the data; identify the field names and data "
   "types; determine which fields are needed for reporting; check for primary, natural and foreign keys; and "
   "quantify what the dataset contains.",
   "The defect classes to look for are: MISSING VALUES (blank, NULL or N/A); DUPLICATED data (the same record "
   "repeated within one dataset); REDUNDANT data (identical data stored in multiple places); INVALID data "
   "(wrong data type, hard-coded values, leading/trailing spaces, non-printable ASCII characters); and "
   "OUTLIERS (values outside the normal distance from the rest of the data).",
   "Cleansing techniques: filter out or impute NULL values (substituting the mean, median or a modelled "
   "estimate); remove duplicates; trim invisible characters and leading/trailing spaces; convert fields to the "
   "correct data type; and decide whether an outlier is an error to correct or a genuine signal to keep.",
   "Tools and techniques: Power Query in Excel, Power BI and Tableau provide built-in profiling; in code, "
   "pandas is used with .isnull().sum() for completeness, .duplicated() for uniqueness, .describe() for the "
   "distribution, and a z-score for outlier detection. "
   "(Slides: Domain 2 — Data Profiling / The Four Defect Classes You Must Find / Data Manipulation Techniques; "
   "Lab 4.)"]),
 ("K2",
  "Descriptive statistics summarise the data you have; inferential statistics let you draw conclusions that "
  "extend beyond the immediate sample. Both underpin every analysis technique in this course.",
  "Explain the statistical principles that underpin data analysis. In your answer, distinguish measures of "
  "central tendency from measures of dispersion, and explain how hypothesis testing and the p-value are used to "
  "decide whether a result is significant.",
  ["CENTRAL TENDENCY describes the centre of a dataset: the MEAN (the sum divided by the count, which uses "
   "every value and is therefore dragged by outliers); the MEDIAN (the middle value when sorted, which resists "
   "outliers); and the MODE (the most frequent value, and the only average that works on categorical data).",
   "DISPERSION describes the spread: the minimum and maximum; the RANGE (max minus min, a first clue that "
   "outliers exist); the VARIANCE (the average squared deviation from the mean); and the STANDARD DEVIATION "
   "(the square root of the variance, expressed in the original units so it can be read directly).",
   "The Z-SCORE standardises a value's distance from the mean, z = (x - x̄) / s. A |z| greater than 3 is the "
   "conventional outlier flag. The EMPIRICAL RULE states that in a normal distribution about 99.74% of values "
   "lie within three standard deviations of the mean.",
   "HYPOTHESIS TESTING states a NULL hypothesis (H0 — there is no relationship between the variables) and an "
   "ALTERNATIVE hypothesis (H1 — a relationship exists) BEFORE the result is seen, then tests them with an "
   "appropriate test such as a t-test (comparing two means) or chi-square (comparing categorical counts).",
   "The P-VALUE is the probability that the observed difference arose by chance. If p is less than 0.05 the "
   "result is statistically significant and H0 is rejected; if p is 0.05 or above, we fail to reject H0. "
   "A TYPE I error is rejecting a true null (a false alarm); a TYPE II error is failing to reject a false null "
   "(a missed effect). Statistical significance does not by itself prove the effect is commercially large. "
   "(Slides: Domain 3 — Central Tendency / Dispersion / Hypothesis Testing; Labs 8 and 9.)"]),
 ("K3",
  "An analysis is only useful once somebody acts on it. The principles of communicating data effectively govern "
  "both which visual you choose and how you present it to a given audience.",
  "Explain the principles of communicating data effectively. Describe how the business question determines the "
  "choice of chart, and how the report or dashboard should be adapted for different audiences.",
  ["The business QUESTION determines the chart, not personal preference. Every question falls into one of five "
   "shapes: composition, comparison, distribution, relationship, or trend over time.",
   "Chart selection: a LINE chart for a trend over time; a BAR or COLUMN chart to compare categories; a PIE "
   "chart for a few parts of one whole; a HISTOGRAM to show the distribution of one variable; a SCATTER PLOT "
   "to test a relationship between two variables; a MAP (dot, filled/choropleth or layered) for geographic "
   "data; a TREEMAP for hierarchical subcategories; and a HEAT MAP to encode magnitude in colour.",
   "Design principles: one purpose per dashboard; consistent colour with one meaning per colour; every axis "
   "labelled and every unit stated; detail hidden behind drillthrough and tooltips rather than crowding the "
   "surface; and a stated refresh date so the reader knows how current the data is.",
   "AUDIENCE adaptation: C-level executives want the big picture — lead with the decision and the single "
   "number that supports it. Management wants performance against target with exceptions and an action. "
   "Technical experts will ask about the method, sample and assumptions, so bring the detail and the caveats. "
   "External stakeholders receive only what may lawfully be shared — aggregate and de-identify first. The "
   "general public needs plain language, one message per visual and a clearly stated source.",
   "The universal rule is to lead with the answer and then the evidence — never make the audience assemble the "
   "conclusion themselves. "
   "(Slides: Domain 3 — Communicating Analysis Results; Domain 4 — Which Chart Answers Which Question / "
   "Dashboard Design Principles; Labs 11 and 12.)"]),
 ("K4",
  "Analysts rarely receive one clean table. Data must be extracted from its source systems and processed into a "
  "single analysis-ready dataset before any statistics are computed.",
  "Explain the procedures used to extract and process data sets. In your answer describe the ETL and ELT "
  "approaches, the SQL join types used to combine data, and the techniques used to transform data once it has "
  "been extracted.",
  ["ETL means Extract, Transform, Load — data is transformed BEFORE it lands in the target, so what arrives in "
   "the data warehouse is already conformed and trustworthy (schema-on-write). ELT means Extract, Load, "
   "Transform — the raw data is landed first, usually in a data lake or lakehouse, and transformed later when "
   "the question is known (schema-on-read). Use ETL when requirements are stable; use ELT when you want to "
   "keep raw data for future, unknown uses.",
   "EXTRACTION connects to the source: a database query (most reliable), an API (pull model polls on a "
   "schedule, push model notifies on change), web scraping (fragile, and requires permission), files such as "
   "CSV/JSON/XML, machine and log data, or public repositories.",
   "LOADING is either a FULL load (all data every run — simple and self-correcting but slow) or a DELTA / "
   "incremental load (only new or changed records — fast, but it depends on a reliable watermark).",
   "SQL JOINS combine datasets on a shared key: INNER JOIN keeps only rows matching in both tables; LEFT OUTER "
   "JOIN keeps all left rows plus matches; RIGHT OUTER JOIN keeps all right rows plus matches; FULL OUTER JOIN "
   "keeps every row from both. An INNER JOIN silently drops unmatched rows, so the record count must always be "
   "reconciled before and after a join.",
   "TRANSFORMATION techniques: recoding values from one form to another; creating DERIVED VARIABLES computed "
   "from existing fields; imputing missing values; aggregating and reducing; transposing or unpivoting; "
   "appending and merging datasets; and parsing strings with delimiters or regular expressions. Query "
   "performance is managed with filtering, indexing, parameterisation, temporary tables and the query "
   "execution plan. "
   "(Slides: Domain 2 — The ETL Pipeline / ETL vs ELT / SQL Join Types / Query Optimisation; Labs 5, 6 and 7.)"]),
]

# ---------------------------------------------------------------- PRACTICAL (ACTIVITY-BASED)
SCENARIO = (
 "You are a data analyst at NovaRetail Pte Ltd, a Singapore retailer with outlets in three regions. Management "
 "has asked you to build the quarterly regional performance review. The data you need sits in three separate "
 "extracts — a customer master, an order transaction file and a regional target file — and the customer file "
 "also carries personal data that must be protected before anything is shared. Complete the five tasks below. "
 "Each task mirrors a hands-on lab you completed in class. For each task, paste your SQL / code and a "
 "screenshot of your output as evidence.")

# (label, criterion, task prompt, box caption, model-answer build steps citing the activity)
BOX_CAP = "Paste your SQL / code and a screenshot of your output in the box below"
PRACTICAL = [
 ("Task 1", "A1",
  "INTEGRATE INFORMATION FROM MULTIPLE DATASETS. You are given three extracts: customers.csv "
  "(customer_id, name, region), orders.csv (order_id, customer_id, amount) and targets.csv (region, target). "
  "Note that one order belongs to a customer who is not in the customer master, and one customer has placed no "
  "orders at all. "
  "Part A — Load all three extracts into SQLite. "
  "Part B — Combine them so that EVERY customer appears in the result, including the customer with no orders, "
  "and enrich each row with that region's target. State which join type you used for each step and why. "
  "Part C — Reconcile the record count: state how many orders exist in the source, how many appear after an "
  "INNER JOIN to customers, and explain the difference. "
  "(Lab 7 — Integrate Multiple Datasets with SQL Joins.)",
  BOX_CAP,
  "Part A — Load the three extracts:\n"
  "sqlite3 integrate.db\n"
  ".mode csv\n"
  ".import customers.csv customers\n"
  ".import orders.csv orders\n"
  ".import targets.csv targets\n\n"
  "Part B — Combine, keeping every customer:\n"
  "SELECT c.region, COUNT(o.order_id) AS orders,\n"
  "       COALESCE(SUM(o.amount),0) AS revenue, t.target\n"
  "FROM customers c\n"
  "LEFT JOIN orders o  ON c.customer_id = o.customer_id\n"
  "JOIN targets t      ON t.region = c.region\n"
  "GROUP BY c.region, t.target;\n\n"
  "Justification: a LEFT JOIN from customers to orders keeps the customer who has no orders (an INNER JOIN "
  "would silently drop them). An INNER JOIN to targets is correct because every region has exactly one target "
  "row. COALESCE(...,0) turns the NULL sum for the order-less region into a reportable zero.\n\n"
  "Part C — Reconciliation: the source holds 5 orders; an INNER JOIN to customers returns only 4. The missing "
  "row is the order whose customer_id has no matching customer master record — an orphan. This is why the "
  "record count must always be reconciled before and after a join. (Lab 7.)"),
 ("Task 2", "A2",
  "CONDUCT A DATA MINING EXERCISE TO UNCOVER TRENDS. Before the data can be analysed you must establish how bad "
  "it is. You are given sales_dirty.csv, which contains blank cells, at least one repeated row and one extreme "
  "value. "
  "Part A — Profile the dataset: report the number of records, the count of missing values per column, and the "
  "number of duplicate rows. "
  "Part B — Detect the outlier using a z-score and state which record is flagged. "
  "Part C — Report both the mean and the median of the spend column, explain which of the two the outlier has "
  "distorted, and state which figure you would report to management and why. "
  "(Lab 4 — Explore a Dirty Dataset; Lab 8 — Descriptive Statistics.)",
  BOX_CAP,
  "Part A — Profile:\n"
  "import pandas as pd\n"
  "d = pd.read_csv('sales_dirty.csv')\n"
  "print(d.shape)                 # record count\n"
  "print(d.isnull().sum())        # missing values per column\n"
  "print(d.duplicated().sum())    # duplicate rows\n\n"
  "Expected: 7 records; 1 null city, 1 null spend, 1 null order_date; 1 duplicate row.\n\n"
  "Part B — Outlier detection with a z-score:\n"
  "s = d['spend'].dropna()\n"
  "z = (s - s.mean()) / s.std()\n"
  "print(d.loc[z[abs(z) > 1.5].index])\n"
  "The record with spend = 99999.00 is flagged as the outlier.\n\n"
  "Part C — Mean vs median:\n"
  "print(d['spend'].mean(), d['spend'].median())\n"
  "The mean (approximately 16,720) is far above the median (approximately 240.50). The MEAN is the statistic "
  "distorted, because it uses every value and is therefore dragged by the extreme 99999.00; the MEDIAN is the "
  "middle value when sorted and resists outliers. Report the MEDIAN to management as the typical transaction "
  "value, and report the outlier separately for investigation. (Labs 4 and 8.)"),
 ("Task 3", "A3",
  "PERFORM DATA ANALYSIS TO DERIVE MEANINGFUL AND ACTIONABLE INSIGHTS. NovaRetail ran an A/B test on a new "
  "checkout page. Group A used the old page and group B the new one; you have the order values for both. "
  "Part A — State the null hypothesis (H0) and the alternative hypothesis (H1) BEFORE running any test. "
  "Part B — Run an appropriate statistical test to compare the two group means, and report the test statistic "
  "and the p-value. "
  "Part C — State your decision against the 0.05 threshold, and write the one-paragraph recommendation you "
  "would give management in plain business language. Identify which error type you would be risking if your "
  "conclusion is wrong. "
  "(Lab 9 — Hypothesis Testing: t-test, p-value and the Two Error Types.)",
  BOX_CAP,
  "Part A — Hypotheses stated first:\n"
  "H0: there is no difference in the mean order value between the old and the new checkout page.\n"
  "H1: the new checkout page has a higher mean order value.\n\n"
  "Part B — Two-sample t-test:\n"
  "import pandas as pd\n"
  "from scipy import stats\n"
  "d = pd.read_csv('abtest.csv')\n"
  "a = d[d.group=='A'].order_value\n"
  "b = d[d.group=='B'].order_value\n"
  "t, p = stats.ttest_ind(a, b)\n"
  "print('t =', t, 'p =', p)\n"
  "Expected: group A averages about 40.75 and group B about 49.5; p is far below 0.05 (approximately 0.000002).\n\n"
  "Part C — Decision and recommendation:\n"
  "Because p < 0.05 we REJECT H0 — the difference is statistically significant and unlikely to be due to "
  "chance. Recommendation: 'The new checkout page increased the average order value by about $8.75 per order "
  "in this test. The result is statistically reliable, so we recommend rolling the new page out to all "
  "customers and monitoring the order value for the next quarter to confirm the lift holds.'\n"
  "Error risk: rejecting H0 when it is in fact true is a TYPE I error (a false positive) — the business would "
  "invest in rolling out a page that actually makes no difference. (Lab 9.)"),
 ("Task 4", "A4",
  "RECOGNISE SEQUENTIAL PATTERNS AND DRAW LINKAGES BETWEEN VARIABLES, AND PROTECT THE DATA. "
  "Part A — Using the monthly marketing dataset (spend, revenue and staff headcount over ten months), compute "
  "the correlation between marketing spend and revenue, fit a regression line, and state the fitted equation "
  "and the R-squared value. "
  "Part B — The dataset shows that staff headcount also correlates strongly with revenue. Explain whether this "
  "means hiring staff CAUSES revenue, identify the confounding variable, and state the principle involved. "
  "Part C — The customer file carries names, NRIC numbers and salaries. Apply a protection technique that "
  "removes the ability to identify an individual while still allowing the departmental salary averages to be "
  "computed, and demonstrate that the analysis still produces the same result. Name the technique you used. "
  "(Lab 10 — Correlation and Regression; Lab 13 — Classify, Mask and De-Identify a Dataset.)",
  BOX_CAP,
  "Part A — Correlation and regression:\n"
  "from scipy import stats\n"
  "r, p = stats.pearsonr(d.spend, d.revenue)\n"
  "lr = stats.linregress(d.spend, d.revenue)\n"
  "print(f'revenue = {lr.slope:.3f} * spend + {lr.intercept:.3f}')\n"
  "print('R-squared =', lr.rvalue**2)\n"
  "Expected: r is approximately 0.999, the fitted line is approximately revenue = 6.6 * spend + 52, and "
  "R-squared is approximately 0.998 — about 99.8% of the variation in revenue is explained by spend.\n\n"
  "Part B — Correlation is not causation:\n"
  "No. Staff headcount correlating with revenue does not mean hiring causes revenue. The CONFOUNDING VARIABLE "
  "is growth over time — the business grew over the ten months, which increased BOTH headcount and revenue "
  "independently. Establishing causation requires a controlled experiment, a plausible mechanism, correct time "
  "ordering and the ruling out of confounders. Observational data alone can never prove it.\n\n"
  "Part C — Pseudonymisation with a surrogate index field:\n"
  "import hashlib\n"
  "d['subject_key'] = d.nric.map(lambda v: hashlib.sha256(v.encode()).hexdigest()[:12])\n"
  "release = d[['subject_key','dept','salary']]\n"
  "print(release.groupby('dept').salary.mean())\n"
  "The technique is PSEUDONYMISATION — the direct identifier (NRIC) is replaced with a non-reversible surrogate "
  "key, and the name and e-mail columns are dropped. The departmental salary averages are identical to those "
  "computed on the original data, proving the analytical value survives the protection. (Masking would hide "
  "part of the value; de-identification would remove the columns outright.) (Labs 10 and 13.)"),
 ("Task 5", "A5",
  "PRESENT ANALYTICS OUTPUTS IN A VISUALLY APPEALING FORMAT TO SUPPORT DECISION-MAKING. Management wants a "
  "one-page quarterly review of regional performance. "
  "Part A — Build a dashboard containing at least four panels: the revenue trend over the three months, a "
  "comparison of total revenue by region, each region's performance against its target, and the relationship "
  "between order count and revenue. For EACH panel, state which chart type you chose and justify it against "
  "the question that panel answers. "
  "Part B — Validate the accuracy of your dashboard using at least three validation techniques, and state the "
  "result of each check. "
  "Part C — State one design choice you made to keep the dashboard readable, and name the audience you "
  "designed it for. "
  "(Lab 11 — Choose the Right Chart; Lab 12 — Build a KPI Dashboard and Validate Its Accuracy.)",
  BOX_CAP,
  "Part A — Four panels with justified chart types:\n"
  "import pandas as pd, matplotlib\n"
  "matplotlib.use('Agg')\n"
  "import matplotlib.pyplot as plt\n"
  "fig, ax = plt.subplots(2, 2, figsize=(12, 7))\n"
  "Panel 1 — LINE chart of revenue by month: the question is 'how is revenue trending?', and the x-axis is "
  "time, so a line chart is correct.\n"
  "Panel 2 — BAR/COLUMN chart of revenue by region: the question is 'which region sells most?', a comparison "
  "of categories, where length is the easiest visual comparison.\n"
  "Panel 3 — BAR chart of percentage of target with a reference line at 100%: the question is 'who is meeting "
  "target?', and the reference line makes pass/fail instantly readable.\n"
  "Panel 4 — SCATTER plot of orders against revenue: the question is 'do more orders mean more revenue?', "
  "which is a relationship between two variables.\n\n"
  "Part B — Validation (at least three techniques):\n"
  "1. RECORD COUNT — confirm the number of source rows equals the number of rows charted (9 region-months).\n"
  "2. RECALCULATION — compute the headline total a second, independent way and confirm it matches:\n"
  "   d.revenue.sum() == sum(d.groupby('region').revenue.sum())  -> total 1129, MATCH.\n"
  "3. CROSS-VALIDATION — check one figure by hand against the code: Central revenue / Central target = 111.9%.\n"
  "   (A peer review or a data audit is also acceptable as a third technique.)\n\n"
  "Part C — Design choice and audience:\n"
  "Consistent colour with one meaning per colour across all four panels, every axis labelled with its unit, "
  "and a reference line rather than extra text on the target panel. Designed for MANAGEMENT — performance "
  "against target with the exceptions visible, and one action implied per panel. (Labs 11 and 12.)"),
]

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → monospace
    code/YAML/command block (indentation preserved); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in code.split("\n"):
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
            rr = b.add_run(ln if ln else " ")
            rr.font.name = "Consolas"; rr.font.size = Pt(9)
            rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
            wt = rr._element.find(qn('w:t'))
            if wt is not None: wt.set(qn('xml:space'), 'preserve')
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    # Pagination is EXPLICIT — two questions to a page on the paper, one model answer to a
    # page in the key. Do not swap this for Word's keepNext/cantSplit: Word pushes an
    # oversized box to the next page, but Google Docs draws the border anyway and prints the
    # question text and the page footer straight THROUGH it. See SKILL.md → Pagination.
    per_page = 1 if answers else 2
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
        if i % per_page == 0 and i < len(WRITTEN):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them.")
        page_break(doc)
    para(doc, "Practical Problem", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    # Practical tasks are long and their boxes are tall, so they get a page each — on the
    # paper AND in the key. Same rule as the WA: the page break is ours, not the renderer's.
    for i, (label, crit, prompt, cap, pts) in enumerate(PRACTICAL, 1):
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=150)
        if i < len(PRACTICAL):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · PP: {len(PRACTICAL)} tasks.")
